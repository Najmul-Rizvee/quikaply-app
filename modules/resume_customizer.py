from openai import OpenAI
from docx import Document
import io

def customize_resume(job_description, api_key):
    client = OpenAI(api_key=api_key)
    prompt = f"Write a tailored resume summary and 3 bullet points for the job described below:\n\n{job_description}"
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=400
    )
    summary = response.choices[0].message.content
    doc = Document()
    doc.add_heading("Customized Resume", 0)
    doc.add_paragraph(summary)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
